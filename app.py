from flask import Flask, render_template, request
import os
import zipfile
from flask import send_file
import re
from flask import Flask, send_file
import openpyxl
import re
from faker import Faker
from docx import Document
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import time

app = Flask(__name__) # 초기화

def make_customer_list():
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    worksheet['A1'] = "이름"
    worksheet['B1'] = "전화번호"
    worksheet['C1'] = "우편번호"
    worksheet['D1'] = "주소"
    worksheet['E1'] = "이메일"

    fake = Faker('ko_KR')


    for row in range(2, 50): #2행부터 시작해 50줄까지 생성
        worksheet.cell(row=row, column=1, value=fake.name())
        worksheet.cell(row=row, column=2, value=fake.phone_number())
        worksheet.cell(row=row, column=3, value=fake.postcode())
        worksheet.cell(row=row, column=4, value=fake.address())
        worksheet.cell(row=row, column=5, value=fake.email())

    workbook.save("customer_list.xlsx")
    return 0


####마스킹 처리####
def process_document():
    dir_path = os.getcwd()
    masking_word_txt = 'word.txt'
    upload_path = "uploads"
    # txt파일에 저장된 특정 단어 읽기
    def read_masking_word(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
            # 공백제거
            return [line.strip() for line in lines]
        
    masking_word = read_masking_word(masking_word_txt)

    # 경로의 doc, docx 파일 찾기 
    for file in os.listdir(dir_path):
        if file.endswith(".docx") or file.endswith(".doc"):
            input_doc = os.path.join(dir_path, file)
            output_doc = os.path.join(upload_path, f"{os.path.splitext(file)[0]}_masking.docx")

            doc = Document(input_doc)

            # doc, docx 파일 특정 단어 변경
            for file_doc in doc.paragraphs:
                for word in masking_word:
                    if word in file_doc.text:
                        # 단어길이만큼 *로 변경
                        masking = '*' * len(word)
                        file_doc.text = file_doc.text.replace(word, masking)
                        
            doc.save(output_doc)
    return(output_doc)


# 한글 파일인 경우 특정 단어가 있으면 마스킹하는 함수
# txt파일을 통해 특정단어 설정
def process_document(file):
    dir_path = 'static'
    masking_word_txt = 'word.txt'
    mst_path = os.path.join(dir_path, masking_word_txt)
    # txt파일에 저장된 특정 단어 읽기
    def read_masking_word(file_path):
        with open(file_path, 'r', encoding='utf-8') as file:
            lines = file.readlines()
            # 공백제거
            return [line.strip() for line in lines]
        
    masking_word = read_masking_word(mst_path)

    # 경로의 doc, docx 파일 찾기 
    upload_path = "uploads"
       
    input_doc = os.path.join(upload_path, file.filename)
    output_doc = os.path.join(upload_path, f"{os.path.splitext(file)[0]}_masking.docx")

    doc = Document(input_doc)

    # doc, docx 파일 특정 단어 변경
    for file_doc in doc.paragraphs:
        for word in masking_word:
            if word in file_doc.text:
                # 단어길이만큼 *로 변경
                masking = '*' * len(word)
                file_doc.text = file_doc.text.replace(word, masking)
                        
    doc.save(output_doc)
    return output_doc

####소스코드 검사####
def make_check_list():
    check_list = []

    with open("check_list.txt", 'r', encoding='utf-8') as f:
        content = f.read()
        lines = [line for line in content.split('\n') if not line.strip().startswith("#")]
        for line in lines:
            if line:
                check_list.extend(line.split(','))
    return check_list


def check_source_code(file):
    check_list = make_check_list()

    # 파일 검사
    filename = file.filename
    # 확장자 확인
    if not filename.endswith(('.py', '.cpp', '.java', '.js', '.kt')):
        # 내용 점검
        content = file.stream.read().decode('utf-8')  # 파일 내용 읽기
        file.stream.seek(0)  # 스트림 포인터를 처음으로 되돌림
        if not any(re.findall(check, content) for check in check_list):
            return True
        
    return False

def make_zip_file(passed_files):
    # 통과한 파일들을 zip 파일로 묶기
    upload_path = "uploads"
    zip_path = os.path.join(upload_path, "compress.zip")
    #os.makedirs(upload_path, exist_ok=True)

    with zipfile.ZipFile(zip_path, 'w') as zip_file:
        for file in passed_files:
            if isinstance(file, str):  # 파일 경로인 경우
                zip_file.write(file, os.path.basename(file))
            else:  # 파일 객체인 경우
                file_path = os.path.join(upload_path, file.filename)
                file.save(file_path)
                zip_file.write(file_path, file.filename)

    compressed_file = "compress.zip"
    return compressed_file


####필터링 한 결과로 메일 재전송 ####
def send_email_with_attachment():
    sender_email = "boanproject1234@naver.com"
    receiver_email = "dbdbtnqls001@naver.com"
    subject = "압축 된 파일입니다"
    body = "압축된 파일"

    zip_file_name = "compress.zip"
    upload_folder = "uploads"

    zip_file_path = os.path.join(upload_folder, zip_file_name)
    with zipfile.ZipFile(zip_file_path, 'w') as zipf:
        for file in os.listdir(upload_folder):
            file_path = os.path.join(upload_folder, file)
            zipf.write(file_path, os.path.basename(file_path))

    smtp_server = "smtp.naver.com"
    smtp_port = 587
    smtp_username = "boanproject1234@naver.com"
    smtp_password = "!Qhdkscjfwj@"

    message = MIMEMultipart()
    message['From'] = sender_email
    message['To'] = receiver_email
    message['Subject'] = subject
    message.attach(MIMEText(body, 'plain'))

    with open(zip_file_path, 'rb') as attachment:
        base = MIMEBase('application', 'zip')
        base.set_payload(attachment.read())
        encoders.encode_base64(base)
        base.add_header('Content-Disposition', 'attachment', filename=zip_file_name)
        message.attach(base)

    with smtplib.SMTP(smtp_server, smtp_port) as server:
        server.starttls()
        server.login(smtp_username, smtp_password)
        server.sendmail(sender_email, receiver_email, message.as_string())

    
    os.remove(zip_file_path)

@app.route('/') # 라우터
def list():
    #make_customer_list()
    return render_template('upload.html')

@app.route('/check_file', methods=["POST"])
def check_file():
    passed_files = []
    files = request.files.getlist("file[]")

    for file in files:
        filename = file.filename
        if filename.endswith('.xlsx'):
            passed_files.append(make_masking_file(file))
        elif filename.endswith(('.docx', 'doc')):
            passed_files.append(process_document(file))
        elif check_source_code(file):
            passed_files.append(file)

    compressed_file = make_zip_file(passed_files)
    compressed_file_path = os.path.join("uploads", compressed_file)
    time.sleep(3)
    send_email_with_attachment()
    return send_file(compressed_file_path, as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)